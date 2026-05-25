"""Convert markdown files to Word documents."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# Combined regex for inline formatting: bold, code, links
INLINE_SPLIT_RE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
LINK_RE = re.compile(r'^\[([^\]]+)\]\(([^)]+)\)$')


def apply_inline_formatting(paragraph, text):
    """Apply bold, code, and link formatting to text within a paragraph."""
    parts = INLINE_SPLIT_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            link_match = LINK_RE.match(part)
            if link_match:
                link_text, url = link_match.group(1), link_match.group(2)
                if url.startswith(('http://', 'https://')):
                    add_hyperlink(paragraph, url, link_text)
                else:
                    # Intra-doc link or relative path — render as plain text
                    paragraph.add_run(link_text)
            else:
                paragraph.add_run(part)


def set_table_borders(table):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    rows = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            # Skip separator lines (|---|---|)
            if not re.match(r'^\|[\s\-:|]+\|$', line):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                rows.append(cells)
        i += 1
    return rows, i


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add as formatted paragraph
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            rows, new_i = parse_table(lines, i)
            if rows:
                # Create table
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                set_table_borders(table)

                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        # Bold header row
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

                doc.add_paragraph()  # Space after table
            i = new_i
            continue

        # Bold text in metadata lines
        if line.startswith('**') and ':**' in line:
            p = doc.add_paragraph()
            apply_inline_formatting(p, line)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text)
            i += 1
            continue

        # Regular paragraph with potential bold/code/link formatting
        p = doc.add_paragraph()
        apply_inline_formatting(p, line.strip())

        i += 1

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def main():
    """Convert all markdown files in the current directory."""
    current_dir = Path(__file__).parent

    md_files = [
        '01-Architecture-Overview.md',
        '02-Staging-Schema.md',
        '03-Warehouse-Schema.md',
        '04-Data-Lineage.md',
        '05-Metrics-Reference.md',
        '06-DYT-Layer2-Schema.md',
        '07-Report-Mapping-Analysis.md',
        'Exasol-Xperts-Narrative.md',
        'Exasol-Xperts-Addendum-A-Market-Sizing.md',
        'Exasol-Xperts-Addendum-B-Shopify-Reporting-Problems.md',
    ]

    for md_file in md_files:
        md_path = current_dir / md_file
        if md_path.exists():
            docx_path = current_dir / md_file.replace('.md', '.docx')
            convert_md_to_docx(md_path, docx_path)
        else:
            print(f"Not found: {md_path}")


if __name__ == '__main__':
    main()
